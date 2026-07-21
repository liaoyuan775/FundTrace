import csv
import hashlib
import json
import shutil
import zipfile
from pathlib import Path

import fitz
import pandas as pd
import rarfile
from docx import Document
from PIL import Image


class UnsafeArchiveError(ValueError):
    pass


class UnsupportedMaterialError(ValueError):
    pass


def safe_extract_zip(
    archive: Path, target: Path, max_files: int = 2000, max_bytes: int = 1073741824
) -> list[Path]:
    target.mkdir(parents=True, exist_ok=True)
    root = target.resolve()
    extracted = []
    total = 0
    with zipfile.ZipFile(archive) as handle:
        infos = handle.infolist()
        if len(infos) > max_files:
            raise UnsafeArchiveError("archive file count exceeds limit")
        for info in infos:
            total += info.file_size
            if total > max_bytes:
                raise UnsafeArchiveError("archive size exceeds limit")
            destination = (target / info.filename).resolve()
            if not destination.is_relative_to(root):
                raise UnsafeArchiveError("archive path traversal detected")
            if info.is_dir():
                destination.mkdir(parents=True, exist_ok=True)
                continue
            destination.parent.mkdir(parents=True, exist_ok=True)
            with handle.open(info) as source, destination.open("wb") as output:
                shutil.copyfileobj(source, output)
            extracted.append(destination)
    return extracted


def safe_extract_rar(
    archive: Path,
    target: Path,
    rar_executable: str = "",
    max_files: int = 2000,
    max_bytes: int = 1073741824,
) -> list[Path]:
    if rar_executable:
        rarfile.UNRAR_TOOL = rar_executable
    target.mkdir(parents=True, exist_ok=True)
    root = target.resolve()
    extracted = []
    total = 0
    try:
        with rarfile.RarFile(archive) as handle:
            infos = handle.infolist()
            if len(infos) > max_files:
                raise UnsafeArchiveError("archive file count exceeds limit")
            if sum(info.file_size for info in infos) > max_bytes:
                raise UnsafeArchiveError("archive size exceeds limit")
            for info in infos:
                destination = (target / info.filename).resolve()
                if not destination.is_relative_to(root):
                    raise UnsafeArchiveError("archive path traversal detected")
                if info.isdir():
                    destination.mkdir(parents=True, exist_ok=True)
                    continue
                destination.parent.mkdir(parents=True, exist_ok=True)
                with handle.open(info) as source, destination.open("wb") as output:
                    while block := source.read(1024 * 1024):
                        total += len(block)
                        if total > max_bytes:
                            raise UnsafeArchiveError("archive size exceeds limit")
                        output.write(block)
                extracted.append(destination)
        return extracted
    except Exception:
        shutil.rmtree(target, ignore_errors=True)
        raise


def extract_material(path: Path, rar_executable: str = "") -> list[dict]:
    suffix = path.suffix.lower()
    chunks = []
    if suffix in {".csv", ".xlsx", ".xls"}:
        sheets = (
            {"Sheet1": pd.read_csv(path)}
            if suffix == ".csv"
            else pd.read_excel(path, sheet_name=None)
        )
        for sheet, frame in sheets.items():
            for index, row in frame.fillna("").iterrows():
                row_data = json.loads(
                    json.dumps(row.to_dict(), ensure_ascii=False, default=str)
                )
                chunks.append(
                    {
                        "text": json.dumps(row_data, ensure_ascii=False, default=str),
                        "row_data": row_data,
                        "sheet_name": sheet,
                        "row_number": int(index) + 2,
                    }
                )
    elif suffix == ".docx":
        doc = Document(path)
        chunks.extend(
            {"text": p.text, "paragraph_number": i + 1}
            for i, p in enumerate(doc.paragraphs)
            if p.text.strip()
        )
        for table_index, table in enumerate(doc.tables):
            row_texts = ["\t".join(cell.text for cell in row.cells) for row in table.rows]
            text = "\n".join(row_texts)
            if text.strip():
                chunks.append(
                    {
                        "text": text,
                        "table_number": table_index + 1,
                        "row_evidence": [
                            {"row_number": row_number, "text": row_text}
                            for row_number, row_text in enumerate(row_texts[1:], start=2)
                        ],
                    }
                )
    elif suffix == ".pdf":
        with fitz.open(path) as doc:
            chunks = [
                {"text": page.get_text("text"), "page_number": i + 1}
                for i, page in enumerate(doc)
            ]
    elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
        with Image.open(path) as image:
            width, height = image.size
        chunks = [
            {
                "text": "",
                "image_path": str(path),
                "requires_vision": True,
                "region": [0, 0, width, height],
            }
        ]
    elif suffix == ".zip":
        files = safe_extract_zip(path, path.parent / (path.stem + "_extracted"))
        for member in files:
            try:
                for item in extract_material(member, rar_executable):
                    chunks.append(
                        {
                            **item,
                            "archive_member_path": str(member.relative_to(path.parent)),
                        }
                    )
            except UnsupportedMaterialError:
                pass
    elif suffix == ".rar":
        if not rar_executable or not Path(rar_executable).exists():
            raise UnsupportedMaterialError("RAR解析器未配置")
        target = path.parent / (path.stem + "_extracted")
        for member in safe_extract_rar(path, target, rar_executable):
            try:
                for item in extract_material(member, rar_executable):
                    chunks.append(
                        {
                            **item,
                            "archive_member_path": str(member.relative_to(target)),
                        }
                    )
            except UnsupportedMaterialError:
                pass
    else:
        raise UnsupportedMaterialError(f"不支持的材料格式: {suffix}")
    return chunks


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()
