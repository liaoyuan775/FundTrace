import argparse
import csv
import json
import os
import shutil
import subprocess
from collections import Counter
from datetime import datetime, timedelta
from pathlib import Path


FILE_SPECS = [
    ("01_甲银行账户流水.csv", 18, 2),
    ("02_乙银行交易明细.xlsx", 18, 2),
    ("03_丙银行电子流水.pdf", 14, 1),
    ("04_丁银行调证材料.docx", 14, 1),
    ("05_手机银行转账记录.png", 9, 1),
    ("06_聚合支付平台明细.xlsx", 13, 2),
    ("07_三方支付账单截图.jpg", 9, 1),
    ("08_公安调取流水汇总.csv", 9, 1),
    ("09_戊银行流水证明.pdf", 8, 2),
    ("10_案件资金流转摘录.docx", 8, 2),
]


def build_truth() -> dict:
    banks = ["甲银行长沙分行", "乙银行深圳分行", "丙银行厦门分行", "丁银行武汉分行", "戊银行南宁分行"]
    regions = ["湖南长沙", "广东深圳", "福建厦门", "湖北武汉", "广西南宁"]
    roles = ["victim"] * 4 + ["level_1"] * 8 + ["level_2"] * 8 + ["level_3"] * 6 + ["exit"] * 6
    names = [
        "受害人甲", "受害人乙", "受害人丙", "受害人丁",
        "张某一", "李某二", "王某三", "赵某四", "周某五", "吴某六", "郑某七", "孙某八",
        "林某九", "陈某十", "杨某十一", "黄某十二", "何某十三", "郭某十四", "罗某十五", "梁某十六",
        "星城商贸", "远航科技", "鸿运百货", "瑞丰电子", "盛汇咨询", "佳信数码",
        "聚合支付商户A", "聚合支付商户B", "数字资产承兑商A", "ATM-长沙解放西路", "POS商户-宏达烟酒", "境外支付通道-XP",
    ]
    entities = [
        {
            "entity_id": f"E{index + 1:02d}",
            "account": f"62{17 + index % 7}{index:015d}",
            "name": names[index],
            "bank": banks[index % len(banks)],
            "region": regions[index % len(regions)],
            "role": roles[index],
        }
        for index in range(32)
    ]
    by_role = {role: [item for item in entities if item["role"] == role] for role in set(roles)}
    transactions = []
    base_time = datetime(2026, 6, 18, 9, 0)

    def add(source: dict, target: dict, stage: str, sequence: int) -> None:
        amount = round(5800 + (sequence % 11) * 3700 + (sequence % 3) * 260.5, 2)
        transaction_time = base_time + timedelta(minutes=sequence * 11 + sequence % 7)
        channel_by_stage = {
            "victim_input": "手机银行", "split": "超级网银", "aggregate": "网上银行",
            "exit": "快捷支付", "return": "银行转账", "third_party": "聚合支付",
        }
        summary_by_stage = {
            "victim_input": "受害人转账", "split": "资金分拆", "aggregate": "资金归集",
            "exit": "消费或取现", "return": "资金回流", "third_party": "三方支付结算",
        }
        index = len(transactions) + 1
        transactions.append({
            "transaction_id": f"FT{index:04d}",
            "transaction_time": transaction_time.isoformat(),
            "serial_number": f"SN20260618{index:06d}",
            "payer_account": source["account"], "payer_name": source["name"], "payer_bank": source["bank"],
            "payee_account": target["account"], "payee_name": target["name"], "payee_bank": target["bank"],
            "debit_credit": "借", "currency": "CNY", "amount": amount,
            "balance_after": round(8000 + (sequence % 13) * 4100.25, 2),
            "channel": channel_by_stage[stage], "summary": summary_by_stage[stage],
            "region": source["region"], "transaction_type": "转账", "stage": stage,
        })

    sequence = 0
    for index in range(16):
        add(by_role["victim"][index % 4], by_role["level_1"][(index * 3) % 8], "victim_input", sequence); sequence += 1
    for index in range(32):
        add(by_role["level_1"][index % 8], by_role["level_2"][(index * 3 + index // 8) % 8], "split", sequence); sequence += 1
    for index in range(32):
        add(by_role["level_2"][index % 8], by_role["level_3"][(index * 5 + index // 8) % 6], "aggregate", sequence); sequence += 1
    for index in range(24):
        add(by_role["level_3"][index % 6], by_role["exit"][(index * 2 + index // 6) % 6], "exit", sequence); sequence += 1
    for index in range(8):
        add(by_role["level_3"][index % 6], by_role["level_1"][(index + 2) % 8], "return", sequence); sequence += 1
    for index in range(8):
        add(by_role["level_2"][index % 8], by_role["exit"][(index + 1) % 6], "third_party", sequence); sequence += 1

    edges = {(item["payer_account"], item["payee_account"]) for item in transactions}
    return {
        "case": {"case_number": "SYN-2026-0717", "name": "多来源资金材料合成测试案"},
        "entities": entities,
        "transactions": transactions,
        "expected_graph": {
            "node_count": len(entities), "edge_count": len(edges), "transaction_count": len(transactions),
            "total_amount": round(sum(item["amount"] for item in transactions), 2),
        },
    }


def build_distribution(transactions: list[dict]) -> list[dict]:
    transaction_ids = [item["transaction_id"] for item in transactions]
    distribution = []
    cursor = 0
    used_duplicate_ids: set[str] = set()
    for file_index, (filename, unique_count, duplicate_count) in enumerate(FILE_SPECS):
        base_ids = transaction_ids[cursor:cursor + unique_count]
        cursor += unique_count
        duplicates = []
        candidate = file_index * 13 + 5
        while len(duplicates) < duplicate_count:
            transaction_id = transaction_ids[candidate % len(transaction_ids)]
            candidate += 17
            if transaction_id in base_ids or transaction_id in used_duplicate_ids:
                continue
            duplicates.append(transaction_id)
            used_duplicate_ids.add(transaction_id)
        distribution.append({
            "filename": filename,
            "format": Path(filename).suffix.lstrip("."),
            "transaction_ids": base_ids + duplicates,
            "unique_source_count": unique_count,
            "duplicate_appearance_count": duplicate_count,
        })
    assert cursor == len(transaction_ids)
    return distribution


def _rows(spec: dict, truth: dict) -> list[dict]:
    by_id = {item["transaction_id"]: item for item in truth["transactions"]}
    return [by_id[transaction_id] for transaction_id in spec["transaction_ids"]]


def _display_row(item: dict) -> list[str]:
    return [
        item["transaction_time"].replace("T", " "), item["serial_number"], item["payer_account"], item["payer_name"],
        item["payee_account"], item["payee_name"], f'{item["amount"]:.2f}', item["channel"], item["summary"],
    ]


def _write_csv(path: Path, rows: list[dict], alternate: bool) -> None:
    if alternate:
        headers = ["发生时间", "银行流水号", "转出账号", "转出户名", "转出银行", "转入账号", "转入户名", "转入银行", "发生额", "账户余额", "支付方式", "用途", "发生地", "币种", "业务类型"]
        fields = ["transaction_time", "serial_number", "payer_account", "payer_name", "payer_bank", "payee_account", "payee_name", "payee_bank", "amount", "balance_after", "channel", "summary", "region", "currency", "transaction_type"]
    else:
        headers = ["交易日期", "交易流水号", "付款账号", "付款户名", "付款行", "收款账号", "收款户名", "收款行", "交易金额", "交易后余额", "交易渠道", "摘要", "地区", "币种", "交易类型"]
        fields = ["transaction_time", "serial_number", "payer_account", "payer_name", "payer_bank", "payee_account", "payee_name", "payee_bank", "amount", "balance_after", "channel", "summary", "region", "currency", "transaction_type"]
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle); writer.writerow(headers)
        for row in rows:
            values = []
            for field in fields:
                value = row[field]
                if field == "transaction_time": value = value.replace("T", " ") if not alternate else value.replace("T", "T")
                if field in {"amount", "balance_after"} and alternate: value = f"¥{value:,.2f}"
                values.append(value)
            writer.writerow(values)


def _write_pdf(path: Path, rows: list[dict], title: str) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Table, TableStyle

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ChineseTitle", parent=styles["Title"], fontName="STSong-Light", fontSize=16, alignment=TA_CENTER)
    document = SimpleDocTemplate(str(path), pagesize=landscape(A4), leftMargin=24, rightMargin=24, topMargin=24, bottomMargin=24)
    elements = [Paragraph(title, title_style)]
    headers = ["交易时间", "流水号", "付款账号", "付款户名", "收款账号", "收款户名", "金额", "渠道", "摘要"]
    for start in range(0, len(rows), 8):
        data = [headers] + [_display_row(item) for item in rows[start:start + 8]]
        table = Table(data, repeatRows=1, colWidths=[95, 100, 105, 55, 105, 70, 58, 60, 70])
        table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"), ("FONTSIZE", (0, 0), (-1, -1), 7.5),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#173B57")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B6C3CC")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF3F6")]),
        ]))
        elements.append(table)
        if start + 8 < len(rows): elements.append(PageBreak())
    document.build(elements)


def _write_docx(path: Path, rows: list[dict], title: str) -> None:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    document = Document(); section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(1.2); section.top_margin = section.bottom_margin = Cm(1.2)
    heading = document.add_paragraph(); heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title); run.bold = True; run.font.name = "Microsoft YaHei"; run.font.size = Pt(15)
    headers = ["交易时间", "流水号", "付款账号", "付款户名", "收款账号", "收款户名", "金额", "渠道", "摘要"]
    table = document.add_table(rows=1, cols=len(headers)); table.style = "Light Shading Accent 1"
    for index, header in enumerate(headers): table.rows[0].cells[index].text = header
    for item in rows:
        cells = table.add_row().cells
        for index, value in enumerate(_display_row(item)): cells[index].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs: run.font.name = "Microsoft YaHei"; run.font.size = Pt(7.5)
    document.add_paragraph("说明：本材料为完全虚构的功能测试数据，不对应任何真实账户或案件。")
    document.save(path)


def _write_image(path: Path, rows: list[dict], title: str, noisy: bool) -> None:
    from PIL import Image, ImageDraw, ImageEnhance, ImageFont

    width, row_height = 2200, 76
    height = 150 + row_height * (len(rows) + 1)
    image = Image.new("RGB", (width, height), "#F4F7FA")
    draw = ImageDraw.Draw(image)
    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    title_font = ImageFont.truetype(str(font_path), 34); header_font = ImageFont.truetype(str(font_path), 22); body_font = ImageFont.truetype(str(font_path), 20)
    draw.rectangle((0, 0, width, 105), fill="#153B5B"); draw.text((42, 28), title, font=title_font, fill="white")
    headers = ["交易时间", "流水号", "付款账号", "付款户名", "收款账号", "收款户名", "金额", "渠道", "摘要"]
    widths = [260, 250, 270, 150, 270, 180, 140, 160, 190]
    x_positions = [30]
    for value in widths[:-1]: x_positions.append(x_positions[-1] + value)
    top = 120
    draw.rectangle((20, top, width - 20, top + row_height), fill="#DDE7EF")
    for index, header in enumerate(headers): draw.text((x_positions[index], top + 22), header, font=header_font, fill="#173B57")
    for row_index, item in enumerate(rows):
        y = top + row_height * (row_index + 1)
        draw.rectangle((20, y, width - 20, y + row_height), fill="#FFFFFF" if row_index % 2 == 0 else "#EEF3F6")
        values = _display_row(item)
        # Group long account identifiers so a vision model can preserve every
        # repeated zero. TransactionRecord removes the visual spaces on ingest.
        values[2] = " ".join(values[2][start:start + 4] for start in range(0, len(values[2]), 4))
        values[4] = " ".join(values[4][start:start + 4] for start in range(0, len(values[4]), 4))
        for column, value in enumerate(values): draw.text((x_positions[column], y + 22), str(value), font=body_font, fill="#1A242C")
    if noisy:
        noise = Image.effect_noise(image.size, 7).convert("RGB")
        image = Image.blend(image, noise, 0.035)
        image = ImageEnhance.Contrast(image).enhance(0.96)
        image.save(path, quality=84, optimize=True)
    else:
        image.save(path)


def generate_base_materials(output_dir: Path) -> None:
    output_dir = Path(output_dir); materials_dir = output_dir / "materials"; oracle_dir = output_dir / "oracle"
    materials_dir.mkdir(parents=True, exist_ok=True); oracle_dir.mkdir(parents=True, exist_ok=True)
    truth = build_truth(); distribution = build_distribution(truth["transactions"])
    appearances = Counter(transaction_id for item in distribution for transaction_id in item["transaction_ids"])
    truth["duplicate_transaction_ids"] = sorted(transaction_id for transaction_id, count in appearances.items() if count > 1)
    (oracle_dir / "ground_truth.json").write_text(json.dumps(truth, ensure_ascii=False, indent=2), "utf-8")
    (oracle_dir / "distribution.json").write_text(json.dumps(distribution, ensure_ascii=False, indent=2), "utf-8")
    with (oracle_dir / "file_manifest.csv").open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle); writer.writerow(["文件名", "格式", "记录出现数", "唯一来源数", "重复出现数"])
        for item in distribution: writer.writerow([item["filename"], item["format"], len(item["transaction_ids"]), item["unique_source_count"], item["duplicate_appearance_count"]])
    for index, spec in enumerate(distribution):
        path = materials_dir / spec["filename"]; rows = _rows(spec, truth); suffix = path.suffix.lower()
        if suffix == ".csv": _write_csv(path, rows, alternate=index > 0)
        elif suffix == ".pdf": _write_pdf(path, rows, path.stem)
        elif suffix == ".docx": _write_docx(path, rows, path.stem)
        elif suffix in {".png", ".jpg"}: _write_image(path, rows, path.stem, suffix == ".jpg")


def generate_xlsx(output_dir: Path, node: str, artifact_tool_module: str = "") -> None:
    if not artifact_tool_module or not Path(artifact_tool_module).exists():
        raise RuntimeError(
            "未找到 @oai/artifact-tool 运行时。请通过 --artifact-tool-module 指定 artifact_tool.mjs 的绝对路径。"
        )
    command = [node, str(Path(__file__).with_name("generate_multisource_xlsx.mjs")), str(output_dir)]
    if artifact_tool_module:
        command.append(artifact_tool_module)
    subprocess.run(command, check=True)


def validate_material_set(output_dir: Path) -> list[str]:
    expected = sorted(filename for filename, _, _ in FILE_SPECS)
    actual = sorted(path.name for path in (Path(output_dir) / "materials").iterdir() if path.is_file())
    if actual != expected:
        raise RuntimeError(f"材料文件集合不完整，期望 {expected}，实际 {actual}")
    return actual


def generate_materials(output_dir: Path, node: str, artifact_tool_module: str = "") -> list[str]:
    generate_base_materials(output_dir)
    generate_xlsx(output_dir, node, artifact_tool_module)
    return validate_material_set(output_dir)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(); parser.add_argument("output_dir", type=Path)
    parser.add_argument("--node", default=os.getenv("FUNDTRACE_NODE") or shutil.which("node") or "")
    parser.add_argument("--artifact-tool-module", default=os.getenv("ARTIFACT_TOOL_MODULE", ""))
    args = parser.parse_args()
    if not args.node:
        raise SystemExit("未找到 Node.js；请通过 --node 指定")
    files = generate_materials(args.output_dir, args.node, args.artifact_tool_module)
    print(json.dumps({"output": str(args.output_dir), "transactions": 120, "appearances": 135, "files": len(files)}, ensure_ascii=False))
